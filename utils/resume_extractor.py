import io


def extract_resume_text(file_bytes: bytes, filename: str) -> dict:
    """
    Extract raw text from uploaded resume (PDF, DOCX, or TXT).
    Returns the text so AI can parse it.
    """
    filename_lower = filename.lower()

    try:
        if filename_lower.endswith(".pdf"):
            text = _extract_from_pdf(file_bytes)
        elif filename_lower.endswith(".docx"):
            text = _extract_from_docx(file_bytes)
        elif filename_lower.endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            return {"error": "Unsupported file format. Use PDF, DOCX, or TXT."}

        if not text or len(text.strip()) < 50:
            return {"error": "Could not extract text from file. Please try a different file."}

        return {"success": True, "text": text.strip()}

    except Exception as e:
        return {"error": f"Failed to read file: {str(e)}"}


def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except ImportError:
        # Fallback to pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except ImportError:
            raise Exception("PDF reading library not found. Run: pip install pymupdf")


def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        raise Exception("DOCX reading library not found. Run: pip install python-docx")
